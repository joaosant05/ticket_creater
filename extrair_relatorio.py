import tkinter as tk
import docx
from tkinter import ttk, messagebox
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt
from pathlib import Path
from mysql.connector import Error
import mimetypes
import magic
import os


class ExtrairRelatorio:
    def __init__(self, root, db, tela_login):
        self.root = root
        self.db = db
        self.tela_login = tela_login
        self.tickets_window = None
        self.tree = None

    def abrir_tela_extrair_relatorio(self):
        if self.tickets_window is not None:
            self.tickets_window.destroy()

        self.tickets_window = tk.Toplevel(self.root)
        self.tickets_window.title("Selecionar Tickets para Relatório")
        self.tickets_window.geometry("790x300")

        mainframe = ttk.Frame(self.tickets_window, padding="10 10 10 10")
        mainframe.grid(column=0, row=0, sticky=(tk.N, tk.W, tk.E, tk.S))
        self.tickets_window.columnconfigure(0, weight=1)
        self.tickets_window.rowconfigure(0, weight=1)

        self.tree = ttk.Treeview(mainframe, columns=("ID", "Tier", "Ambiente", "Frequência", "Usuário", "Navegador", "Organização", "Marca", "Data Criação"), show='headings')

        columns = ["ID", "Tier", "Ambiente", "Frequência", "Usuário", "Navegador", "Organização", "Marca", "Data Criação"]
        col_widths = {
            "ID": 20,
            "Tier": 40,
            "Ambiente": 80,
            "Frequência": 80,
            "Usuário": 150,
            "Navegador": 80,
            "Organização": 80,
            "Marca": 80,
            "Data Criação": 140
        }

        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[col], anchor=tk.W, stretch=tk.NO)

        self.tree.grid(column=0, row=1, sticky=(tk.N, tk.W, tk.E, tk.S))

        scrollbar = ttk.Scrollbar(mainframe, orient=tk.VERTICAL, command=self.tree.yview)
        scrollbar.grid(column=1, row=1, sticky=(tk.N, tk.S))
        self.tree.configure(yscrollcommand=scrollbar.set)

        ttk.Button(mainframe, text="Voltar", command=self.voltar_para_login).grid(column=0, row=2, sticky=tk.W, padx=5, pady=10)
        ttk.Button(mainframe, text="Gerar Relatório", command=self.gerar_relatorio).grid(column=0, row=2, sticky=tk.E, padx=5, pady=10)

        self.atualizar_tickets()

    def atualizar_tickets(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        connection = self.db.conectar_banco()
        if connection is not None:
            try:
                cursor = connection.cursor(dictionary=True)
                query = "SELECT * FROM relatorios"
                cursor.execute(query)
                tickets = cursor.fetchall()

                for ticket in tickets:
                    self.tree.insert("", tk.END, values=(
                        ticket.get("id"),
                        ticket.get("tier"),
                        ticket.get("ambiente"),
                        ticket.get("frequencia"),
                        ticket.get("usuario"),
                        ticket.get("navegador"),
                        ticket.get("organizacao"),
                        ticket.get("marca"),
                        ticket.get("created_at"),
                    ))

            except Error as e:
                messagebox.showerror("Erro", f"Erro ao buscar tickets do MySQL: {e}")
            finally:
                if connection.is_connected():
                    cursor.close()
                    connection.close()

    def atualizar_tickets(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        connection = self.db.conectar_banco()
        if connection is not None:
            try:
                cursor = connection.cursor(dictionary=True)
                query = "SELECT * FROM relatorios"
                cursor.execute(query)
                tickets = cursor.fetchall()

                for ticket in tickets:
                    self.tree.insert("", tk.END, values=(
                        ticket.get("id"),
                        ticket.get("tier"),
                        ticket.get("ambiente"),
                        ticket.get("frequencia"),
                        ticket.get("usuario"),
                        ticket.get("navegador"),
                        ticket.get("organizacao"),
                        ticket.get("marca"),
                        ticket.get("created_at"),
                    ))

            except Error as e:
                messagebox.showerror("Erro", f"Erro ao buscar tickets do MySQL: {e}")
            finally:
                if connection.is_connected():
                    cursor.close()
                    connection.close()

    def gerar_relatorio(self):
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("Aviso", "Por favor, selecione pelo menos um ticket.")
            return

        document = Document()
        document.add_heading('Relatório de Tickets', 0)

        # Criar a pasta para salvar os anexos
        anexos_pasta = self.criar_pasta_anexos()

        for item in selected_items:
            item_values = self.tree.item(item, "values")
            ticket_id = item_values[0]

            connection = self.db.conectar_banco()
            if connection is not None:
                try:
                    cursor = connection.cursor(dictionary=True)
                    query = "SELECT * FROM relatorios WHERE id = %s"
                    cursor.execute(query, (ticket_id,))
                    ticket = cursor.fetchone()

                    document.add_heading(f"Ticket ID: {ticket_id}", level=1)

                    self.adicionar_paragrafo_formatado(document, "Tier:", ticket.get('tier'))
                    self.adicionar_paragrafo_formatado(document, "Ambiente:", ticket.get('ambiente'))
                    self.adicionar_paragrafo_formatado(document, "Frequência:", ticket.get('frequencia'))
                    self.adicionar_paragrafo_formatado(document, "Usuário:", ticket.get('usuario'))
                    self.adicionar_paragrafo_formatado(document, "Navegador:", ticket.get('navegador'))
                    self.adicionar_paragrafo_formatado(document, "Organização:", ticket.get('organizacao'))
                    self.adicionar_paragrafo_formatado(document, "Marca:", ticket.get('marca'))

                    par_log = document.add_paragraph()
                    run_log = par_log.add_run("Log:")
                    run_log.bold = True
                    run_log.font.size = Pt(12)
                    par_log.add_run(f"\n{ticket.get('log')}")

                    self.adicionar_paragrafo_formatado(document, "Como Reproduzir:", ticket.get('como_reproduzir'))
                    self.adicionar_paragrafo_formatado(document, "O Que Acontece:", ticket.get('o_que_acontece'))
                    self.adicionar_paragrafo_formatado(document, "O Que Deveria Acontecer:", ticket.get('o_que_deveria_acontecer'))

                    data_criacao = ticket.get('created_at')
                    if data_criacao:
                        data_formatada = data_criacao.strftime("%d/%m/%Y - %H:%M")
                        self.adicionar_paragrafo_formatado(document, "Data Criação:", data_formatada)

                    query_anexos = "SELECT * FROM anexos WHERE ticket_id = %s"
                    cursor.execute(query_anexos, (ticket_id,))
                    anexos = cursor.fetchall()

                    for anexo in anexos:
                        self.processar_anexo(document, anexo, anexos_pasta)

                except Error as e:
                    messagebox.showerror("Erro", f"Erro ao buscar ticket do MySQL: {e}")
                finally:
                    if connection.is_connected():
                        cursor.close()
                        connection.close()

        self.salvar_arquivo(document, "relatorio_tickets.docx")

    def voltar_para_login(self):
        self.tickets_window.destroy()
        self.root.deiconify()
        self.tela_login.criar_tela_login()

    def adicionar_paragrafo_formatado(self, document, titulo, texto):
        par = document.add_paragraph()
        run_titulo = par.add_run(titulo)
        run_titulo.bold = True
        run_titulo.font.size = Pt(12)
        par.add_run(f" {texto}")

    def processar_anexo(self, document, anexo, anexos_pasta):
        blob_data = anexo['anexo']
        if blob_data:
            mime_type = self.deduzir_mime_type(blob_data)
            if mime_type:
                arquivo_path = self.salvar_blob_como_arquivo(blob_data, mime_type, anexos_pasta)
                if mime_type.startswith('image'):
                    try:
                        document.add_picture(str(arquivo_path), width=Inches(2))  # Converte Path para string
                    except Exception as e:
                        messagebox.showerror("Erro", f"Erro ao inserir imagem: {e}")
                elif mime_type.startswith('video'):
                    paragraph = document.add_paragraph(f"(anexo vídeo: ")
                    self.add_hyperlink(paragraph, arquivo_path.as_uri(), "Clique para abrir o vídeo")
                    paragraph.add_run(")")
                elif mime_type == 'application/pdf':
                    paragraph = document.add_paragraph(f"(anexo PDF: ")
                    self.add_hyperlink(paragraph, arquivo_path.as_uri(), "Clique para abrir o PDF")
                    paragraph.add_run(")")
                else:
                    paragraph = document.add_paragraph(f"Anexo desconhecido ({mime_type}): ")
                    self.add_hyperlink(paragraph, arquivo_path.as_uri(), "Clique para abrir")
            else:
                messagebox.showwarning("Aviso", "Tipo MIME não reconhecido.")
        else:
            messagebox.showwarning("Aviso", "Anexo está vazio.")

    def salvar_arquivo(self, document, filename="relatorio_tickets.docx"):
        desktop_path = Path(os.path.join(os.path.expanduser("~"), "Desktop", filename))
        onedrive_path = Path(os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop", filename))
        downloads_path = Path(os.path.join(os.path.expanduser("~"), "Downloads", filename))

        if desktop_path.parent.exists():
            document.save(desktop_path)
            print(f"Arquivo salvo no Desktop: {desktop_path}")
        elif onedrive_path.parent.exists():
            document.save(onedrive_path)
            print(f"Arquivo salvo no OneDrive: {onedrive_path}")
        elif downloads_path.parent.exists():
            document.save(downloads_path)
            print(f"Arquivo salvo na pasta de Downloads: {downloads_path}")
        else:
            messagebox.showerror("Erro", "Não foi possível salvar o arquivo em nenhum diretório padrão.")

    def deduzir_mime_type(self, blob_data):
        mime = magic.Magic(mime=True)
        return mime.from_buffer(blob_data)

    def salvar_blob_como_arquivo(self, blob_data, mime_type, anexos_pasta):
        extensao = mimetypes.guess_extension(mime_type)
        arquivo_path = anexos_pasta / f"anexo_{int(Path().stat().st_mtime)}{extensao}"  # Nome único
        with open(arquivo_path, 'wb') as file:
            file.write(blob_data)
        return arquivo_path

    def criar_pasta_anexos(self):
        # Primeiro, tentamos criar a pasta no Desktop
        desktop_path = Path(os.path.join(os.path.expanduser("~"), "Desktop"))
        onedrive_desktop_path = Path(os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop"))

        anexos_pasta = desktop_path / "Anexos_Tickets"
        
        if desktop_path.exists():
            # Se o Desktop normal existir, criamos a pasta lá
            anexos_pasta.mkdir(exist_ok=True)
            return anexos_pasta
        elif onedrive_desktop_path.exists():
            # Se o Desktop estiver no OneDrive, criamos a pasta no Desktop do OneDrive
            anexos_pasta = onedrive_desktop_path / "Anexos_Tickets"
            anexos_pasta.mkdir(exist_ok=True)
            return anexos_pasta
        else:
            # Caso nenhum dos dois caminhos exista, lançamos um erro
            raise FileNotFoundError("Não foi possível encontrar o caminho do Desktop ou do OneDrive.")

    def add_hyperlink(self, paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')

        new_run.append(r_pr)
        new_run.text = text

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
